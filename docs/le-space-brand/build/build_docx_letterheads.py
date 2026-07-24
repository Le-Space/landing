#!/usr/bin/env python3
"""Le-Space company letter templates (no personal name), DE/EN — python-docx."""
import os
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "briefkopf")
LOGO = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "logo", "png", "le-space-logo-horizontal-light@2048.png")

INK = RGBColor(0x14, 0x1B, 0x2E)
GREY = RGBColor(0x4A, 0x54, 0x68)
LIGHT = RGBColor(0x6B, 0x76, 0x90)
PLACEHOLDER = RGBColor(0x9A, 0xA3, 0xB5)

COMPANY = "Le-Space"
STREET = "Lichtenberg 44"
CITY = "D-84307 Eggenfelden"
COUNTRY = "Germany"
PHONE = "+49 174 9 89 19 49"
MAIL = "nico.krause@le-space.de"
WEB = "le-space.de"

STR = {
    "de": dict(claim="Der Local-First Peer-to-Peer Stack",
               date="Eggenfelden, [Datum]", subject="Betreff: …",
               recip="[Empfängeradresse]",
               salute="Sehr geehrte Damen und Herren,", body="[Ihr Text]",
               close="Mit freundlichen Grüßen", sig="[Name]"),
    "en": dict(claim="The Local-First Peer-to-Peer Stack",
               date="Eggenfelden, [Date]", subject="Subject: …",
               recip="[Recipient address]",
               salute="Dear Sir or Madam,", body="[Your text]",
               close="Kind regards", sig="[Name]"),
}

def set_border(par, edge, color="D8DEE9", size=6):
    pPr = par._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr'); pPr.append(pBdr)
    el = OxmlElement(f'w:{edge}')
    el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(size))
    el.set(qn('w:space'), '4'); el.set(qn('w:color'), color)
    pBdr.append(el)

def run(par, text, font="Arial", size=11, color=INK, bold=False):
    r = par.add_run(text)
    r.font.name = font; r.font.size = Pt(size)
    r.font.color.rgb = color; r.bold = bold
    r._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    return r

def spacing(par, before=0, after=0, line=None):
    pf = par.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    if line: pf.line_spacing = line

def build(lang):
    s = STR[lang]
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210); sec.page_height = Mm(297)
    sec.top_margin = Mm(12); sec.bottom_margin = Mm(28)
    sec.left_margin = Mm(25); sec.right_margin = Mm(25)
    sec.header_distance = Mm(10); sec.footer_distance = Mm(10)

    # ---- header ----
    hdr = sec.header
    hdr.paragraphs[0].text = ""
    tbl = hdr.add_table(1, 2, Mm(160))
    tbl.autofit = False
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none'); el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), 'auto')
        borders.append(el)
    tblPr.append(borders)
    tbl.columns[0].width = Mm(80); tbl.columns[1].width = Mm(80)
    c1, c2 = tbl.rows[0].cells
    c1.width = Mm(80); c2.width = Mm(80)
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c1.paragraphs[0]; spacing(p, after=2)
    p.add_run().add_picture(LOGO, width=Mm(52))
    p = c1.add_paragraph(); spacing(p)
    run(p, s["claim"], "Arial", 7, LIGHT)
    # right company block
    p = c2.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; spacing(p)
    run(p, COMPANY, "Arial", 9, INK, bold=True)
    for t in (STREET, CITY, COUNTRY):
        p = c2.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; spacing(p)
        run(p, t, "Arial", 7.5, GREY)
    for t in (PHONE, MAIL, WEB):
        p = c2.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; spacing(p)
        run(p, t, "Consolas", 6.5, GREY)
    rulep = hdr.add_paragraph(); spacing(rulep, before=6)
    set_border(rulep, "bottom")

    # ---- footer ----
    ftr = sec.footer
    fp = ftr.paragraphs[0]; fp.text = ""; spacing(fp, after=3)
    set_border(fp, "top")
    p = ftr.add_paragraph(); spacing(p)
    run(p, f"{COMPANY} — {s['claim']}", "Consolas", 6, LIGHT)
    p = ftr.add_paragraph(); spacing(p)
    run(p, f"{STREET} · {CITY} · {COUNTRY} · {PHONE} · {MAIL} · {WEB}", "Consolas", 6, LIGHT)

    # ---- body ----
    # DIN sender line + recipient block
    p = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    spacing(p, before=24, after=2)
    r = run(p, f"{COMPANY} · {STREET} · {CITY} · {COUNTRY}", "Arial", 6.5, LIGHT)
    set_border(p, "bottom", size=4)
    p = doc.add_paragraph(); spacing(p, before=6, after=48)
    run(p, s["recip"], "Arial", 10, PLACEHOLDER)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; spacing(p, after=24)
    run(p, s["date"], "Arial", 10, GREY)
    p = doc.add_paragraph(); spacing(p, after=12)
    run(p, s["subject"], "Arial", 11, INK, bold=True)
    p = doc.add_paragraph(); spacing(p, after=12)
    run(p, s["salute"], "Arial", 11, INK)
    p = doc.add_paragraph(); spacing(p, after=36)
    run(p, s["body"], "Arial", 11, GREY)
    p = doc.add_paragraph(); spacing(p, after=24)
    run(p, s["close"], "Arial", 11, INK)
    p = doc.add_paragraph()
    run(p, s["sig"], "Arial", 11, PLACEHOLDER)

    path = os.path.join(OUT, f"brief-vorlage-le-space-{lang}.docx")
    doc.save(path)
    print("docx:", os.path.basename(path))

for lang in ("de", "en"):
    build(lang)
print("DONE")
